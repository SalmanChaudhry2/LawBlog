from flask import Flask, render_template, request, redirect, url_for, session, send_from_directory, jsonify, g, send_file
from docx import Document
from openai import AzureOpenAI
from dotenv import load_dotenv
import os
import time
from datetime import datetime, timedelta
import requests
import markdown
from bs4 import BeautifulSoup
from flask_session import Session
import json
import sqlite3
from io import BytesIO
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

load_dotenv()

app = Flask(__name__, static_folder='static')
app.secret_key = os.getenv('FLASK_SECRET_KEY')
app.config['UPLOAD_FOLDER'] = 'static/generated'
app.config['DATABASE'] = os.path.join(app.instance_path, 'app.db')

# Session configuration
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=1)
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
Session(app)

# Database functions
def get_db():
    if 'db' not in g:
        g.db = sqlite3.connect(app.config['DATABASE'])
        g.db.row_factory = sqlite3.Row
    return g.db

def close_db(e=None):
    db = g.pop('db', None)
    if db is not None:
        db.close()

def init_db():
    os.makedirs(app.instance_path, exist_ok=True)
    with app.app_context():
        db = get_db()
        cursor = db.cursor()
        
        # Create users table
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            firm TEXT,
            location TEXT,
            lawyer_name TEXT,
            state TEXT,
            keywords TEXT    
        )
        ''')
        
        # Create tones table
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS tones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            description TEXT NOT NULL,
            UNIQUE(user_id, name),
            FOREIGN KEY(user_id) REFERENCES users(id)
        )
        ''')
        
        # Insert default users if they don't exist
        cursor.execute("SELECT * FROM users WHERE username IN ('admin', 'memberhub')")
        existing_users = cursor.fetchall()
        existing_usernames = [user['username'] for user in existing_users]
        
        if 'admin' not in existing_usernames:
            cursor.execute('''
            INSERT INTO users (username, email, password, firm, location, lawyer_name, state, keywords)
            VALUES (?, ?, ?, ?, ?, ?, ?,?)
            ''', (
                'admin', 
                'admin@lawfirm.com', 
                'password123', 
                'Legal Partners', 
                'New York', 
                'John', 
                'NY',
                'Personal Family Lawyer,Estate Planning Firm, Life & Legacy Planning'
            ))
        
        if 'memberhub' not in existing_usernames:
            cursor.execute('''
            INSERT INTO users (username, email, password, firm, location, lawyer_name, state, keywords)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                'memberhub', 
                'memberhub@newlawbusinessmodel.com', 
                'memberhub123', 
                'New Law Business Model', 
                'Global', 
                'Member Hub', 
                'CA',
                'Trusted estate planning firm, Asset protection services, Estate planning attorney'
            ))
        
        db.commit()

# Initialize database
with app.app_context():
    init_db()

# Add context processor to inject current year into all templates
@app.context_processor
def inject_year():
    return {'now': datetime.now()}

class Config:
    ARTICLES_DIR = "articles"
    GENERATED_DIR = "generated"
    os.makedirs(ARTICLES_DIR, exist_ok=True)
    os.makedirs(GENERATED_DIR, exist_ok=True)

class AzureServices:
    def __init__(self):
        self.text_client = AzureOpenAI(
            api_key=os.getenv("AZURE_OPENAI_KEY"),
            api_version="2024-02-15-preview",
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
        )
        
        self.conversations = {}

    def rewrite_content(self, original_text, tone, tone_description, keywords, firm_name, location, lawyer_name, city, state, planning_session_name="Life & Legacy Planning Session"):
        response = self.text_client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=[
                {"role": "system", "content": f"""
                    You are a legal blog post rewriter. There should be At least 30% changes from original. Rewrite the article following these strict guidelines:
                    SEO REQUIREMENTS:
                    1. Must include these elements within the first 150 words:
                       - Primary keywords: {keywords}
                       - Firm name: {firm_name}
                       - City-state of firm: {location}
                       - Lawyer name: {lawyer_name}
                       - City-state of Lawyer: {city}, {state}
                    2. Incorporate naturally - don't just list them
                    
                    TONE REQUIREMENTS:
                    1. Primary Tone: {tone}
                    2. Tone Description: {tone_description}
                    3. Consistency: Maintain this tone throughout the entire article
                    
                    SPECIAL BRANDING REQUIREMENTS:
                    - Avoid transactional language like "investing in" which are not aligned with the Personal Family Lawyer® brand tone
                    - Instead use phrases like:
                        * "work with us to choose a plan that works to keep your loved ones out of court and out of conflict"
                        * "create a plan that protects what matters most"
                        * "develop a comprehensive approach to safeguarding your family's future"
                        * "put a plan in place that ensures your wishes are honored"
                        * "create a plan that grows with your family and ensures lasting peace of mind"
                    - Emphasize the ongoing relationship and family protection aspects rather than transactional terms
                    - Use the term "{planning_session_name}" when referencing to planning sessions.

                    CONTENT GUIDELINES:
                    DO's:
                    1. Use active voice
                    2. Structure with 5 sections: introduction, 3 subheadings, and conclusion with call-to-action
                    3. Keep length between 1000-1200 words
                    4. Use transition sentences between sections
                    5. Conclusion should be brief (1-2 sentences) with clear call-to-action
                    6. Include 1-2 bulleted lists in the entire article
                    7. Balance paragraphs and lists appropriately
                    8. Write in a {tone} tone
                    9. Include these keywords naturally: {keywords}
                    10. Mention {firm_name} in {location} where relevant
                    11. Firm name is {firm_name} and location is {location}
                    12 Lawyer name is {lawyer_name} and location is {city}, {state}
                    
                    DON'Ts:
                    1. Avoid legal jargon or complex language (keep it high-school level)
                    2. No passive voice
                    3. Don't use lists without context
                    4. Limit metaphors
                    5. Don't make conclusion too long
                    6. Don't include more than 5 sources
                    7. Don't exceed 1200 words
                    8. Don't use more than 3 lists
                    
                    CTA REQUIREMENTS:
                    1. MUST use the exact phrase "15-minute Discovery Call" (never "consultation" or "consult")
                    2. Standard format: "Schedule your complimentary 15-minute Discovery Call with {firm_name} today"
                    3. Include a clear call-to-action like "Click here to schedule" or "Book your Discovery Call now"
                    4. Never offer to answer questions or provide consultation during this call

                    STYLE GUIDE UPDATES:
                    1. LANGUAGE PREFERENCE:
                    - Use "loved ones" instead of "family" in all cases EXCEPT when:
                        * Referring specifically to legal family members (spouse, children, parents)
                        * Discussing family law matters specifically related to spouse, children, parents
                        * The context explicitly requires "family" (e.g., "family business")
                    - Preferred phrases:
                        * "protect your loved ones"
                        * "ensure your loved ones are cared for"
                        * "keep your loved ones out of court"
                        * "provide for your loved ones"

                    Formatting Requirements:
                    # Main Title
                    ## Subheading 1
                    ### Sub-subheading (if needed)
                    **Bold important terms**
                    - Bullet points when appropriate
                    [Link text](URL) for references
                    
                    The article must be valuable, engaging, and optimized for both readers and search engines.
                """},
                {"role": "user", "content": original_text}
            ],
            temperature=0.7,
        )
        return response.choices[0].message.content

    def edit_content(self, session_id, user_message, current_content=None):
        if session_id not in self.conversations:
            self.conversations[session_id] = [
                {"role": "system", "content": """
                    You are a legal blog post editor. When the user requests changes:
                    1. Make ONLY the requested changes
                    2. Return the COMPLETE updated blog (not just updated part) in markdown format
                    3. Don't include any commentary or explanations
                    4. Preserve all formatting and structure
                """}
            ]
        
        if current_content:
            self.conversations[session_id].append(
                {"role": "assistant", "content": current_content}
            )
        
        self.conversations[session_id].append(
            {"role": "user", "content": user_message}
        )
        
        response = self.text_client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=self.conversations[session_id],
            temperature=0.5
        )
        
        ai_response = response.choices[0].message.content
        self.conversations[session_id].append(
            {"role": "assistant", "content": ai_response}
        )
        
        return ai_response
    
class ImageGenerator:
    def __init__(self):
        self.image_client = AzureOpenAI(
            api_key=os.getenv("AZURE_DALLE_KEY"),
            api_version="2024-02-01",
            azure_endpoint=os.getenv("AZURE_DALLE_ENDPOINT")
        )
        self.text_client = AzureOpenAI(
            api_key=os.getenv("AZURE_OPENAI_KEY"),
            api_version="2024-02-15-preview",
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
        )

    def generate_image(self, text_prompt):
        try:
            safe_prompt = self._get_safe_image_prompt(text_prompt)
            
            response = self.image_client.images.generate(
                model=os.getenv("AZURE_DALLE_DEPLOYMENT"),
                prompt=safe_prompt,
                size="1024x1024",
                quality="standard",
                n=1,
            )
            image_url = response.data[0].url
            os.makedirs(os.path.join(app.static_folder, 'generated'), exist_ok=True)
            
            timestamp = int(time.time())
            image_filename = f"image_{timestamp}.png"
            image_path = os.path.join(app.static_folder, 'generated', image_filename)
            
            response = requests.get(image_url)
            with open(image_path, 'wb') as f:
                f.write(response.content)
            
            return image_filename
            
        except Exception as e:
            print(f"Image generation failed: {e}")
            return None
        
    def _get_safe_image_prompt(self, text_prompt):
        response = self.text_client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=[
                {"role": "system", "content": """
                    You are a creative prompt engineer for legal blog images. Create safe and professional image prompts that:
                    1. Are directly relevant to the blog content
                    2. Be 'unique to the blog's content', not generic or reusable for any legal article
                    3. Reflect the main topic, themes, or message of the blog post
                    4. Focus on modern, visually appealing representations
                    5. Must pass Azure content filters
                    6. Avoids sensitive content
                    The prompt should be detailed and specific, including:
                        - Main subject
                        - Style description
                        - Color palette
                        - Composition notes
                        - Mood/tone
                    - Is based on this blog content:
                """},
                {"role": "user", "content": text_prompt[:1000]}
            ],
            temperature=1
        )
        return response.choices[0].message.content

class FileManager:
    @staticmethod
    def list_articles():
        """
        List all DOCX files in the articles directory
        Returns:
            List of article filenames
        """
        articles = [f for f in os.listdir(Config.ARTICLES_DIR) if f.endswith('.docx')]
        return articles
    
    @staticmethod
    def get_article_metadata():
        """
        Read and parse the metadata.json file
        Returns:
            Dictionary of article metadata
        """
        metadata_path = os.path.join(Config.ARTICLES_DIR, 'metadata.json')
        try:
            with open(metadata_path, 'r', encoding='utf-8') as f:
                content = f.read()
                metadata = json.loads(content)
                # Convert list to dictionary for easier lookup
                result = {article['filename']: article for article in metadata['articles']}
                return result
        except (FileNotFoundError, json.JSONDecodeError, KeyError) as e:
            print(f"Error reading metadata: {str(e)}")
            return {}
    
    @staticmethod
    def read_docx(filename):
        """
        Read content from a DOCX file
        Args:
            filename: Name of the DOCX file
        Returns:
            Extracted text content
        """
        doc = Document(os.path.join(Config.ARTICLES_DIR, filename))
        return "\n".join([para.text for para in doc.paragraphs])
    
    @staticmethod
    def save_content(content):
        """
        Save generated content to a file
        Args:
            content: Content to save
        Returns:
            Filename of the saved content
        """
        filename = f"blog_{int(time.time())}.txt"
        path = os.path.join(Config.GENERATED_DIR, filename)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return filename

    @staticmethod
    def generate_formatted_docx(content, title="Legal Blog"):
        """Generate DOCX with exact formatting from markdown"""
        doc = Document()

        # Custom styles (can be modified)
        styles = {
            'h1': {'font_size': 16, 'bold': True, 'color': RGBColor(0, 32, 96)},
            'h2': {'font_size': 14, 'bold': True, 'color': RGBColor(0, 64, 128)},
            'h3': {'font_size': 12, 'bold': True, 'italic': True},
            'bold': {'bold': True},
            'normal': {'font_size': 11}
        }
        
        def apply_style(run, style):
            """Helper function to apply formatting"""
            run.font.size = Pt(style.get('font_size', 11))
            run.font.bold = style.get('bold', False)
            run.font.italic = style.get('italic', False)
            if 'color' in style:
                run.font.color.rgb = style['color']
        
        # Process markdown content line by line
        lines = content.split('\n')
        for line in lines:

            if line.replace('-', '').strip() == '' and len(line) >= 3:
                continue

            # Detect formatting
            if line.startswith('# '):  # H1
                para = doc.add_heading(level=1)
                run = para.add_run(line[2:].strip())
                apply_style(run, styles['h1'])
                
            elif line.startswith('## '):  # H2
                para = doc.add_heading(level=2)
                run = para.add_run(line[3:].strip())
                apply_style(run, styles['h2'])
                
            elif line.startswith('### '):  # H3
                para = doc.add_heading(level=3)
                run = para.add_run(line[4:].strip())
                apply_style(run, styles['h3'])
                
            elif '**' in line:  # Bold text
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.+?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        apply_style(run, styles['bold'])
                    else:
                        para.add_run(part)
            
            else:  # Normal paragraph
                para = doc.add_paragraph()
                run = para.add_run(line)
                apply_style(run, styles['normal'])
        # Collect all empty paragraphs
        empty_paragraphs = [p for p in doc.paragraphs if not p.text.strip()]

        # Remove each empty paragraph from the document
        for p in empty_paragraphs:
            p._element.getparent().remove(p._element)

        # Save to bytes buffer
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

class UserSession:
    @staticmethod
    def register(email, password, firm, location, lawyer_name, state, keywords=""):
        db = get_db()
        username = email.split('@')[0].lower()
        try:
            cursor = db.cursor()
            cursor.execute('''
            INSERT INTO users (username, email, password, firm, location, lawyer_name, state, keywords)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (username, email, password, firm, location, lawyer_name, state, keywords))
            db.commit()
            return True
        except sqlite3.IntegrityError:
            return False

    @staticmethod
    def login(email, password):
        db = get_db()
        username = email.split('@')[0].lower()
        user = db.execute(
            'SELECT * FROM users WHERE username = ?', 
            (username,)
        ).fetchone()
        
        if user and user['password'] == password:
            # Get user's custom tones
            tones = db.execute(
                'SELECT name, description FROM tones WHERE user_id = ?',
                (user['id'],)
            ).fetchall()
            
            session['user'] = {
                'id': user['id'],
                'username': user['username'],
                'email': user['email'],
                'firm': user['firm'],
                'location': user['location'],
                'lawyer_name': user['lawyer_name'],
                'state': user['state'],
                'keywords': user['keywords'],
                'custom_tones': [dict(tone) for tone in tones]
            }
            return True
        return False

    @staticmethod
    def update_profile(username, firm, location, lawyer_name, state, keywords):
        db = get_db()
        try:
            cursor = db.cursor()
            cursor.execute('''
            UPDATE users 
            SET firm = ?, location = ?, lawyer_name = ?, state = ?, keywords = ?
            WHERE username = ?
            ''', (firm, location, lawyer_name, state, keywords, username))
            db.commit()
            
            # Update session if this is the current user
            if 'user' in session and session['user']['username'] == username:
                session['user'].update({
                    'firm': firm,
                    'location': location,
                    'lawyer_name': lawyer_name,
                    'state': state,
                    'keywords': keywords
                })
                session.modified = True
            return True
        except sqlite3.Error:
            return False
        
    @staticmethod
    def add_custom_tone(user_id, tone_name, tone_description):
        db = get_db()
        try:
            cursor = db.cursor()
            cursor.execute('''
            INSERT INTO tones (user_id, name, description)
            VALUES (?, ?, ?)
            ''', (user_id, tone_name, tone_description))
            db.commit()
            
            # Update session if this is the current user
            if 'user' in session and session['user']['id'] == user_id:
                session['user']['custom_tones'].append({
                    'name': tone_name,
                    'description': tone_description
                })
                session.modified = True
            return True
        except sqlite3.IntegrityError:
            return False
    
    @staticmethod
    def get_custom_tones(user_id):
        db = get_db()
        tones = db.execute(
            'SELECT name, description FROM tones WHERE user_id = ?',
            (user_id,)
        ).fetchall()
        return [dict(tone) for tone in tones]
    
    @staticmethod
    def get_current_user():
        return session.get('user')

azure_services = AzureServices()
image_generator = ImageGenerator()

@app.template_filter('markdown')
def markdown_filter(text):
    html = markdown.markdown(text)
    soup = BeautifulSoup(html, 'html.parser')
    return str(soup)

@app.route('/')
def home():
    if not UserSession.get_current_user():
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        firm = request.form['firm']
        location = request.form['location']
        lawyer_name = request.form['lawyer_name']
        state = request.form['state']
        keywords = request.form.get('keywords', '')
        if UserSession.register(email, password, firm, location, lawyer_name, state, keywords):
            # Auto-login after registration
            UserSession.login(email, password)
            return redirect(url_for('dashboard'))
        
        return render_template('register.html', error="Email already registered")
    
    return render_template('register.html')

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        firm = request.form['firm']
        location = request.form['location']
        lawyer_name = request.form['lawyer_name']
        state = request.form['state']
        keywords = request.form.get('keywords', '')
        
        if UserSession.update_profile(user['username'], firm, location, lawyer_name, state, keywords):
            session['user']['firm'] = firm
            session['user']['location'] = location
            session['user']['lawyer_name'] = lawyer_name
            session['user']['state'] = state
            session['user']['keywords'] = keywords
            session.modified = True
            
            return redirect(url_for('dashboard'))
        
        return render_template('profile.html', error="Update failed", user=session['user'])
    
    return render_template('profile.html', user=session['user'])

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if UserSession.login(request.form['email'], request.form['password']):
            return redirect(url_for('dashboard'))
        return render_template('login.html', error="Invalid credentials")
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    
    # Get articles and their metadata
    articles = FileManager.list_articles()
    metadata = FileManager.get_article_metadata()
    
    # Combine standard tones with user's custom tones
    standard_tones = [
        ('Professional', 'Formal and business-like tone suitable for corporate audiences'),
        ('Friendly', 'Warm and approachable tone that builds rapport with readers'),
        ('Educational', 'Clear and informative tone designed to explain concepts')
    ]
    
    custom_tones = user.get('custom_tones', [])
    all_tones = standard_tones + [(t['name'], t['description']) for t in custom_tones]
    
    # Convert to the format expected by the template
    tone_options = [t[0] for t in all_tones]
    tone_descriptions = {t[0]: t[1] for t in all_tones}
    
    return render_template('dashboard.html', 
                         user=user,
                         username=user['username'],
                         articles=articles,
                         metadata=metadata,
                         tone_options=tone_options,
                         tone_descriptions=tone_descriptions,
                         user_keywords=user.get('keywords', ''))

@app.route('/add_tone', methods=['POST'])
def add_tone():
    user = UserSession.get_current_user()
    if not user:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    data = request.get_json() if request.is_json else request.form
    tone_name = data.get('tone_name', '').strip()
    tone_description = data.get('tone_description', '').strip()
    
    if not tone_name:
        return jsonify({'success': False, 'error': 'Tone name is required'}), 400
    
    if UserSession.add_custom_tone(user['id'], tone_name, tone_description):
        return jsonify({
            'success': True,
            'tone_name': tone_name,
            'tone_description': tone_description
        })
    
    return jsonify({'success': False, 'error': 'Tone with this name already exists'}), 400

@app.route('/select/<article>', methods=['GET', 'POST'])
def select_article(article):
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    if request.method == 'POST':
        tone = request.form.get('tone')
        tone_description = request.form.get('toneDescription')
        custom_tone = request.form.get('customToneName')

        if tone == 'custom' and custom_tone:
            tone = custom_tone
            
        keywords = request.form.get('keywords', '')
        firm = request.form.get('firm', '')
        location = request.form.get('location', '')
        lawyer_name = user.get('lawyer_name', '')
        city = user.get('location', '')
        state = user.get('state', '')
        planning_session_name = request.form.get('planning_session_name','') 
        if not planning_session_name:
            planning_session_name="Life & Legacy Planning Session"

        # Generate the blog post with the selected tone
        blog_content = azure_services.rewrite_content(
            FileManager.read_docx(article),
            tone,
            tone_description,
            keywords,
            firm,
            location,
            lawyer_name,
            city,
            state,
            planning_session_name
        )
        
        # Save the generated content to a file
        filename = FileManager.save_content(blog_content)
        
        # Set up the session data for the review page (without image initially)
        session['current_post'] = {
            'original': article,
            'content': blog_content,
            'image': None,  # Image will be generated later when requested
            'created': datetime.now().strftime("%Y-%m-%d %H:%M"),
            'tone': tone,
            'filename': filename
        }
        
        # Initialize chat history
        session['chat_history'] = [{
            'role': 'assistant',
            'content': blog_content,
            'content_is_blog': True,
            'timestamp': datetime.now().strftime("%H:%M:%S")
        }]
        
        # Generate a unique session ID for the chat
        session['session_id'] = os.urandom(16).hex()
        
        return redirect(url_for('review'))
    
    tone_options = [
        'Professional',
        'Friendly',
        'Educational'
    ]
    
    tone_descriptions = {
        'Professional': 'Formal and business-like tone suitable for corporate audiences',
        'Friendly': 'Warm and approachable tone that builds rapport with readers',
        'Educational': 'Clear and informative tone designed to explain concepts'
    }
    
    return render_template('select.html',
                         article_name=article,
                         tone_options=tone_options,
                         tone_descriptions=tone_descriptions,
                         firm=firm,
                         location=location)

@app.route('/use_version', methods=['POST'])
def use_version():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    selected_content = request.form['content']
    
    session['current_post']['content'] = selected_content
    session.modified = True
    
    return redirect(url_for('finalize'))

@app.route('/finalize')
def finalize():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    post = session['current_post']
    filename = FileManager.save_content(post['content'])
    image_url = url_for('static', filename=f'generated/{post["image"]}') if post.get('image') else None
    
    return render_template('finalize.html', 
                         post=post,
                         filename=filename,
                         image_url=image_url)

@app.route('/review', methods=['GET', 'POST'])
def review():
    # Check if we have a filename parameter but no current_post in session
    filename = request.args.get('filename')
    if filename and 'current_post' not in session:
        # Try to load the content from the file
        try:
            with open(os.path.join(Config.GENERATED_DIR, filename), 'r', encoding='utf-8') as f:
                content = f.read()
                
            # Set up the session data
            session['current_post'] = {
                'content': content,
                'filename': filename,
                'created': datetime.now().strftime("%Y-%m-%d %H:%M")
            }
            
            # Initialize chat history
            session['chat_history'] = [{
                'role': 'assistant',
                'content': content,
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            }]
            
            # Generate a unique session ID for the chat
            session['session_id'] = os.urandom(16).hex()
        except Exception as e:
            print(f"Error loading file: {e}")
            return redirect(url_for('dashboard'))
    
    # If we still don't have current_post in session, redirect to dashboard
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    post = session['current_post']
    
    if 'session_id' not in session:
        session['session_id'] = os.urandom(16).hex()
    
    if 'chat_history' not in session:
        session['chat_history'] = [{
            'role': 'assistant',
            'content': post['content'],
            'content_is_blog': True,
            'timestamp': datetime.now().strftime("%H:%M:%S")
        }]
    
    if request.method == 'POST':
        if 'edit_message' in request.form:  # Chat-style editing
            user_message = request.form['edit_message']
            
            current_content = next(
                (msg['content'] for msg in reversed(session['chat_history']) 
                 if msg['content_is_blog']),
                post['content']
            )
            
            edited_content = azure_services.edit_content(
                session['session_id'],
                user_message,
                current_content
            )
            
            session['chat_history'].append({
                'role': 'user',
                'content': user_message,
                'content_is_blog': False,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })
            session['chat_history'].append({
                'role': 'assistant',
                'content': edited_content,
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })
            
            post['content'] = edited_content
            session['current_post'] = post
            
        elif 'content' in request.form:  # Manual editing
            post['content'] = request.form['content']
            session['current_post'] = post
            session['chat_history'].append({
                'role': 'assistant',
                'content': post['content'],
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })

        session.modified = True
        return redirect(url_for('review'))
    
    # Save the current content to a file and get the filename
    if 'filename' not in post:
        filename = FileManager.save_content(post['content'])
        post['filename'] = filename
        session['current_post'] = post
    
    image_url = url_for('static', filename=f'generated/{post["image"]}') if post.get('image') else None
    
    return render_template('review.html', 
                         post=post,
                         chat_history=session['chat_history'],
                         image_url=image_url)

@app.route('/save_changes', methods=['POST'])
def save_changes():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    edited_content = request.form.get('content', '')
    
    session['current_post']['content'] = edited_content
    
    if 'chat_history' not in session:
        session['chat_history'] = []
    
    session['chat_history'].append({
        'role': 'system',
        'content': 'User saved manual changes',
        'content_is_blog': False,
        'timestamp': datetime.now().strftime("%H:%M:%S")
    })
    
    session.modified = True
    return redirect(url_for('finalize'))

@app.route('/download/<filename>')
def download(filename):
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    try:
        # Read generated content
        filepath = os.path.join(Config.GENERATED_DIR, filename)
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Get title
        title = session['current_post'].get('original', 'Legal Blog').replace('.docx', '')
        
        # Generate formatted DOCX
        docx_file = FileManager.generate_formatted_docx(content, title)
        
        return send_file(
            docx_file,
            as_attachment=True,
            download_name=f"{title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"DOCX generation failed: {e}")
        return redirect(url_for('review'))

@app.route('/generate_image')
def generate_image():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    # Generate image based on current content
    image_filename = image_generator.generate_image(session['current_post']['content'])
    
    if image_filename:
        session['current_post']['image'] = image_filename
        session.modified = True
    
    return redirect(url_for('review'))
    
@app.teardown_appcontext
def teardown_db(exception):
    close_db()

if __name__ == '__main__':
    app.run(debug=True)